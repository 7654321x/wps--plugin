"""Read-only validation of the saved WPS one-click formatting result.

This validator deliberately reads the persisted DOCX rather than trusting an
in-memory WPS property assignment.  Its output only contains paragraph
indices, expected roles and failed property names: no document text or paths.
"""

from __future__ import annotations

import argparse
import hashlib
import json
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from lxml import etree


NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
ALIGNMENT = {"left": "left", "center": "center", "right": "right", "justify": "both", "distributed": "distribute"}


def attr(node, name):
    element = node.find("w:" + name, NS) if node is not None else None
    return {key.rsplit("}", 1)[-1]: value for key, value in element.attrib.items()} if element is not None else {}


def saved_paragraphs(path):
    with ZipFile(path) as archive:
        root = etree.fromstring(archive.read("word/document.xml"))
    rows = []
    for paragraph in root.xpath(".//w:body/w:p", namespaces=NS):
        ppr = paragraph.find("w:pPr", NS)
        rpr = paragraph.find(".//w:r/w:rPr", NS)
        rows.append({
            "alignment": attr(ppr, "jc").get("val"),
            "indent": attr(ppr, "ind"),
            "spacing": attr(ppr, "spacing"),
            "outline": attr(ppr, "outlineLvl").get("val"),
            "fonts": attr(rpr, "rFonts"),
            "size": attr(rpr, "sz").get("val"),
            "bold": attr(rpr, "b"),
        })
    return rows


def verify(document_path, profile_path, expectations_path):
    profile = json.loads(profile_path.read_text(encoding="utf-8"))
    expectations = json.loads(expectations_path.read_text(encoding="utf-8"))
    document = Document(document_path)
    saved = saved_paragraphs(document_path)
    text_parts = [paragraph.text for paragraph in document.paragraphs]
    failures = []
    for item in expectations["paragraphs"]:
        index = item["index"]
        role = item["role"]
        style = profile["styles"][role]
        actual = saved[index] if index < len(saved) else None
        if actual is None:
            failures.append({"index": index, "role": role, "property": "paragraph"})
            continue
        expected_outline = 9 if style["outline_level"] == 10 else style["outline_level"] - 1
        expected = {
            "font": style["east_asia_font_name"], "latin_font": style["latin_font_name"],
            "size": str(int(style["font_size_pt"] * 2)), "bold": style["bold"],
            "alignment": ALIGNMENT[style["alignment"]], "first_line_chars": str(int(style["first_line_indent_chars"] * 100)),
            "left_chars": str(int(style["left_indent_chars"] * 100)), "right_chars": str(int(style["right_indent_chars"] * 100)),
            "line": str(int(style["line_spacing_pt"] * 20)), "outline": str(expected_outline),
        }
        observed = {
            "font": actual["fonts"].get("eastAsia"), "latin_font": actual["fonts"].get("ascii"),
            "size": actual["size"], "bold": actual["bold"].get("val", "1") != "0",
            "alignment": actual["alignment"], "first_line_chars": actual["indent"].get("firstLineChars"),
            "left_chars": actual["indent"].get("leftChars"), "right_chars": actual["indent"].get("rightChars"),
            "line": actual["spacing"].get("line"), "outline": actual["outline"],
        }
        # DOCX has no run properties for a truly empty paragraph.  Its
        # alignment, indentation, spacing and outline are still persisted and
        # validated; a font assertion would have no stored target to read.
        if not document.paragraphs[index].text:
            for name in ("font", "latin_font", "size", "bold"):
                expected.pop(name)
        for name, value in expected.items():
            if observed[name] != value:
                failures.append({"index": index, "role": role, "property": name})
    return {
        "status": "PASS" if not failures else "FAIL",
        "body_sha256": hashlib.sha256("\x1f".join(text_parts).encode("utf-8")).hexdigest(),
        "paragraph_count": len(text_parts),
        "section_count": len(document.sections),
        "failures": failures,
    }


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("--profile", type=Path, required=True)
    parser.add_argument("--expectations", type=Path, required=True)
    parser.add_argument("--out", type=Path)
    args = parser.parse_args()
    result = verify(args.docx, args.profile, args.expectations)
    rendered = json.dumps(result, ensure_ascii=False, indent=2) + "\n"
    if args.out:
        args.out.parent.mkdir(parents=True, exist_ok=True)
        args.out.write_text(rendered, encoding="utf-8")
    else:
        print(rendered, end="")
    raise SystemExit(0 if result["status"] == "PASS" else 1)


if __name__ == "__main__":
    main()

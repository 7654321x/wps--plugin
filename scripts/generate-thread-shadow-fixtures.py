from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


COUNTS = (20, 200, 1000)


def paragraph_text(index: int) -> str:
    if index % 37 == 0:
        return ""
    if index % 29 == 0:
        return f"脱敏线程快照超长段落 {index:04d}：" + "中文与 Latin text 仅用于验证本地批次读取。" * 12
    return f"脱敏线程快照测试段落 {index:04d}：中文与 Latin text。"


def create_fixture(path: Path, paragraph_count: int) -> None:
    document = Document()
    document.core_properties.title = f"{paragraph_count} 段线程快照脱敏测试"
    document.core_properties.subject = "Docxtool WPS snapshot shadow fixture"
    document.core_properties.author = "Docxtool Test Fixture"
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    for index in range(1, paragraph_count + 1):
        paragraph = document.add_paragraph(paragraph_text(index))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def main() -> int:
    parser = argparse.ArgumentParser(description="生成 WPS 后台线程快照脱敏测试文档")
    parser.add_argument("--output-dir", type=Path, required=True)
    args = parser.parse_args()
    for count in COUNTS:
        path = args.output_dir / f"snapshot-shadow-{count}.docx"
        create_fixture(path, count)
        actual = len(Document(path).paragraphs)
        if actual != count:
            raise RuntimeError(f"FIXTURE_PARAGRAPH_COUNT_MISMATCH:{count}:{actual}")
        print(f"FIXTURE_READY {path} paragraphs={actual}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

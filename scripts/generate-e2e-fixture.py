"""Generate anonymous one-click-format fixtures with intentionally plain text."""

import argparse
from datetime import datetime, timezone
from pathlib import Path
from tempfile import NamedTemporaryFile
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from lxml import etree


FIXTURE_LINES = (
    "Docxtool 一键排版自动验收",
    "主标题测试段落",
    "一、一级标题测试段落",
    "（一）二级标题测试段落",
    "1.三级标题测试段落",
    "（1）四级标题测试段落",
    "普通正文测试段落：这是脱敏示例文字，用于验证正文格式是否写入。",
    "中英文混排正文：Docxtool WPS 123 ABC 与中文混排。",
    "称呼测试段落：各位同志：",
    "附件说明测试段落：附件：1.测试材料",
    "附件正文标题测试段落",
    "附件正文测试段落：仅用于自动验收。",
    "落款署名测试段落：测试单位",
    "落款日期测试段落：2026年7月30日",
    "重复段落测试",
    "重复段落测试",
    "重复段落测试",
    "",
    "已有错误缩进段落：此行从零缩进起始。",
    "已有错误行距段落：此行用于覆盖错误行距。",
    "已有错误字体段落：ABC 123 中文。",
    "左对齐",
    "居中",
    "右对齐",
    "两端对齐：这是用于观察两端对齐效果的脱敏示例文字。",
    "分散对齐：这是用于观察分散对齐效果的脱敏示例文字。",
)
USER_COMMENT_TEXT = "用户保留批注：此批注必须在预览和一键排版后保持。"


def remove_grids(document):
    for section in document.sections:
        grid = section._sectPr.find(qn("w:docGrid"))
        if grid is not None:
            section._sectPr.remove(grid)


def build_document():
    document = Document()
    comment_target = None
    for index, text in enumerate(FIXTURE_LINES):
        paragraph = document.add_paragraph()
        # A multi-run paragraph is still intentionally direct-format free.
        if index == 7:
            paragraph.add_run("中英文混排正文：Docxtool ")
            paragraph.add_run("WPS 123 ABC")
            paragraph.add_run(" 与中文混排。")
        else:
            paragraph.add_run(text)
        if index == 6:
            comment_target = paragraph
    if comment_target is not None and hasattr(document, "add_comment"):
        document.add_comment(
            runs=comment_target.runs,
            text=USER_COMMENT_TEXT,
            author="验收用户",
            initials="YS",
        )
    document.add_section()
    document.add_paragraph("纵向分节测试：该节用于验证多节页面设置。")
    landscape = document.add_section()
    landscape.orientation = WD_ORIENT.LANDSCAPE
    landscape.page_width, landscape.page_height = landscape.page_height, landscape.page_width
    document.add_paragraph("横向分节测试：该节必须在一键排版后保持横向。")
    final_section = document.add_section()
    final_section.orientation = WD_ORIENT.PORTRAIT
    if final_section.page_width > final_section.page_height:
        final_section.page_width, final_section.page_height = final_section.page_height, final_section.page_width
    document.add_paragraph("纵横混合多节测试：最后一节恢复纵向。")
    remove_grids(document)
    return document


def _inject_comment_ooxml(path, paragraph_index=6):
    """Add one synthetic user comment when python-docx lacks comments support."""
    with ZipFile(path) as source:
        entries = {name: source.read(name) for name in source.namelist()}

    document = etree.fromstring(entries["word/document.xml"])
    namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    paragraphs = document.xpath("//w:body/w:p", namespaces=namespaces)
    paragraph = paragraphs[paragraph_index]
    start = etree.Element(qn("w:commentRangeStart"))
    start.set(qn("w:id"), "0")
    end = etree.Element(qn("w:commentRangeEnd"))
    end.set(qn("w:id"), "0")
    reference_run = etree.Element(qn("w:r"))
    reference = etree.SubElement(reference_run, qn("w:commentReference"))
    reference.set(qn("w:id"), "0")
    paragraph.insert(0, start)
    paragraph.append(end)
    paragraph.append(reference_run)
    entries["word/document.xml"] = etree.tostring(document, xml_declaration=True, encoding="UTF-8", standalone=True)

    w_ns = namespaces["w"]
    comments = etree.Element(qn("w:comments"), nsmap={"w": w_ns})
    comment = etree.SubElement(comments, qn("w:comment"))
    comment.set(qn("w:id"), "0")
    comment.set(qn("w:author"), "验收用户")
    comment.set(qn("w:initials"), "YS")
    comment.set(qn("w:date"), datetime.now(timezone.utc).isoformat())
    comment_paragraph = etree.SubElement(comment, qn("w:p"))
    run = etree.SubElement(comment_paragraph, qn("w:r"))
    text = etree.SubElement(run, qn("w:t"))
    text.text = USER_COMMENT_TEXT
    entries["word/comments.xml"] = etree.tostring(comments, xml_declaration=True, encoding="UTF-8", standalone=True)

    relationships = etree.fromstring(entries["word/_rels/document.xml.rels"])
    rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    ids = {item.get("Id", "") for item in relationships}
    number = 1
    while f"rId{number}" in ids:
        number += 1
    relation = etree.SubElement(relationships, f"{{{rel_ns}}}Relationship")
    relation.set("Id", f"rId{number}")
    relation.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments")
    relation.set("Target", "comments.xml")
    entries["word/_rels/document.xml.rels"] = etree.tostring(relationships, xml_declaration=True, encoding="UTF-8", standalone=True)

    content_types = etree.fromstring(entries["[Content_Types].xml"])
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    override = etree.SubElement(content_types, f"{{{ct_ns}}}Override")
    override.set("PartName", "/word/comments.xml")
    override.set("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml")
    entries["[Content_Types].xml"] = etree.tostring(content_types, xml_declaration=True, encoding="UTF-8", standalone=True)

    with NamedTemporaryFile(delete=False, suffix=".docx", dir=path.parent) as temporary:
        temporary_path = Path(temporary.name)
    try:
        with ZipFile(temporary_path, "w", ZIP_DEFLATED) as target:
            for name, value in entries.items():
                target.writestr(name, value)
        temporary_path.replace(path)
    finally:
        temporary_path.unlink(missing_ok=True)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--output-dir", type=Path, default=Path(__file__).resolve().parents[1] / "tests" / "fixtures")
    args = parser.parse_args()
    args.output_dir.mkdir(parents=True, exist_ok=True)
    document = build_document()
    # These are plain-text baselines. The WPS host writes the after-format copy.
    for name in ("wps-e2e-baseline.docx", "01-before-format.docx", "02-rollback-test.docx", "03-after-one-click-format.docx"):
        output = args.output_dir / name
        document.save(output)
        if not hasattr(document, "add_comment"):
            _inject_comment_ooxml(output)


if __name__ == "__main__":
    main()

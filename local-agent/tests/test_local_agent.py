import io
import json
import hashlib
import sys
from pathlib import Path
from wsgiref.util import setup_testing_defaults

ROOT = Path(__file__).resolve().parents[1]
WPS_ROOT = ROOT.parent
sys.path.insert(0, str(ROOT / "src"))
sys.path.insert(0, str(WPS_ROOT / "command-service" / "src"))

from docxtool_local_agent.app import create_app  # noqa: E402
from docxtool_local_agent.e2e import (  # noqa: E402
    guard_test_document,
    record_diagnostics,
    record_result,
    run_readonly_chain,
)


def _call(app, method, path, body=None, token="test-token", origin=""):
    environ = {}
    setup_testing_defaults(environ)
    environ["REQUEST_METHOD"] = method
    environ["PATH_INFO"] = path
    environ["HTTP_X_DOCXTOOL_SESSION"] = token
    if origin:
        environ["HTTP_ORIGIN"] = origin
    data = json.dumps(body or {}).encode("utf-8")
    environ["CONTENT_LENGTH"] = str(len(data))
    environ["wsgi.input"] = io.BytesIO(data)
    captured = {}
    response = b"".join(app(environ, lambda status, headers: captured.update(status=status, headers=dict(headers))))
    return captured["status"], json.loads(response.decode("utf-8")), captured["headers"]


def test_health_and_version_are_text_free():
    app = create_app("test-token")
    status, health, _ = _call(app, "GET", "/v1/health")
    assert status == "200 OK"
    assert health == {"ok": True}
    status, version, _ = _call(app, "GET", "/v1/version")
    assert status == "200 OK"
    assert "source_path" not in json.dumps(version)
    status, handshake, _ = _call(app, "GET", "/v1/handshake")
    assert status == "200 OK"
    assert handshake["ok"] is True
    assert handshake["host_text_contract_version"] == "host-text-v1"
    serialized_handshake = json.dumps(handshake)
    assert "DocxTool local handshake" not in serialized_handshake
    assert "source_path" not in serialized_handshake


def test_recognition_returns_a_redacted_sdk_plan_and_keeps_input_unchanged(tmp_path):
    from docx import Document

    source = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("严禁进入响应的正文")
    document.save(source)
    original = source.read_bytes()
    status, payload, _ = _call(create_app("test-token"), "POST", "/v1/recognize", {
        "source_path": str(source),
        "host_snapshot": {
            "host_type": "wps-test",
            "paragraphs": [{"host_paragraph_index": 0, "raw_text": "严禁进入响应的正文"}],
        },
    })
    assert status == "200 OK"
    assert source.read_bytes() == original
    serialized = json.dumps(payload, ensure_ascii=False)
    assert payload["data"]["blocks"][0]["recognized_text"] == "严禁进入响应的正文"
    assert str(source) not in serialized
    assert payload["data"]["blocks"][0]["text_sha256"]
    assert payload["data"]["blocks"][0]["locator_verified"] is True
    assert payload["binding"]["blocks"][0]["binding_status"] == "confirmed"
    assert payload["binding"]["blocks"][0]["host_raw_start_utf16"] == 0


def test_recognition_requires_a_local_host_snapshot(tmp_path):
    from docx import Document

    source = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("脱敏文本")
    document.save(source)

    status, payload, _ = _call(create_app("test-token"), "POST", "/v1/recognize", {
        "source_path": str(source),
    })

    assert status == "400 Bad Request"
    assert payload["error"]["code"] == "INVALID_REQUEST"


def test_single_local_service_builds_commands_on_the_same_port():
    request = {
        "schema_version": "1.0", "request_id": "request-00000001",
        "recognition_result": {
            "schema_version": "1.2", "recognition_engine_version": "3.0",
            "document_id": "doc-1", "document_revision": "rev-1",
            "source_sha256": "a" * 64, "document_mode": "normal",
            "document_mode_confidence": 1, "paragraphs": [],
        },
        "profile_id": "default", "profile_version": "1.0",
        "client_capabilities": {"schema_version": "1.0", "capabilities": []},
        "product_version": "0.1.0", "authorization_scope": "classified-offline",
    }
    status, payload, _ = _call(create_app("test-token"), "POST", "/v1/commands", request)
    assert status == "200 OK"
    assert payload["request_id"] == request["request_id"]
    assert payload["commands"] == []


def test_cors_only_allows_the_fixed_taskpane_origin_and_preflight_headers():
    app = create_app("test-token")
    status, _, headers = _call(app, "OPTIONS", "/v1/e2e/session", origin="http://127.0.0.1:3889")
    assert status == "204 No Content"
    assert headers["Access-Control-Allow-Origin"] == "http://127.0.0.1:3889"
    assert "POST" in headers["Access-Control-Allow-Methods"]
    assert "Authorization" in headers["Access-Control-Allow-Headers"]
    _, _, rejected = _call(app, "OPTIONS", "/v1/e2e/session", origin="http://127.0.0.1:3999")
    assert "Access-Control-Allow-Origin" not in rejected


def test_e2e_result_is_redacted_and_test_guard_requires_the_session_copy(tmp_path):
    runtime = tmp_path / "e2e"
    session_id = "a" * 32
    runtime.mkdir()
    (runtime / "current.json").write_text(json.dumps({
        "session_id": session_id, "test_results": {}, "overall_status": "REAL_WPS_E2E_NOT_RUN",
    }), encoding="utf-8")
    session_dir = runtime / session_id
    session_dir.mkdir()
    working = session_dir / "test-working-copy.docx"
    working.write_bytes(b"safe fixture bytes")
    (session_dir / "test-document.json").write_text(json.dumps({
        "working_file": "test-working-copy.docx",
        "working_sha256": hashlib.sha256(working.read_bytes()).hexdigest(),
        "is_fixture_baseline": False,
    }), encoding="utf-8")
    assert guard_test_document(runtime, session_id, str(working)) == {"ok": True}
    assert guard_test_document(runtime, session_id, str(tmp_path / "other.docx"))["error_code"] == "E2E_TEST_DOCUMENT_REQUIRED"
    assert record_result(runtime, {
        "session_id": session_id, "stage": "runtime_probe", "status": "PASS", "error_code": "",
    }) == {"ok": True}
    assert record_result(runtime, {
        "session_id": session_id, "stage": "paragraph.set_alignment", "status": "PASS", "error_code": "",
    }) == {"ok": True}
    assert (runtime.parent / "capabilities" / session_id / "paragraph.set_alignment.json").is_file()
    assert record_diagnostics(runtime, {"session_id": session_id, "diagnostics": [{
        "check_id": "LOCAL_AGENT_HEALTH", "group": "本地服务", "status": "PASS",
        "error_code": "", "summary": "健康", "duration_ms": 1, "dependencies": [],
    }]}) == {"ok": True, "count": 1}
    with __import__("pytest").raises(ValueError, match="REDACTION"):
        record_result(runtime, {
            "session_id": session_id, "stage": "x", "status": "PASS", "error_code": "", "text": "forbidden",
        })
    assert "safe fixture" not in (runtime / "current.json").read_text(encoding="utf-8")


def test_readonly_chain_uses_verified_host_bindings_and_excludes_text(tmp_path):
    from docx import Document

    runtime = tmp_path / "e2e"
    session_id = "b" * 32
    runtime.mkdir()
    (runtime / "current.json").write_text(json.dumps({
        "session_id": session_id, "plugin_version": "1.0", "test_results": {},
        "overall_status": "REAL_WPS_E2E_NOT_RUN",
    }), encoding="utf-8")
    session_dir = runtime / session_id
    session_dir.mkdir()
    working = session_dir / "test-working-copy.docx"
    document = Document()
    document.add_paragraph("一、脱敏标题。")
    document.add_paragraph("脱敏正文内容。")
    document.save(working)
    (session_dir / "test-document.json").write_text(json.dumps({
        "working_file": working.name,
        "working_sha256": hashlib.sha256(working.read_bytes()).hexdigest(),
        "is_fixture_baseline": False,
    }), encoding="utf-8")

    result = run_readonly_chain(runtime, session_id, include_plan=True)

    assert result["ok"] is True
    assert result["command_count"] >= 1
    assert result["anchors"]
    assert all(item["formatting_disposition"] in {"apply", "skip"} for item in result["anchors"])
    serialized = json.dumps(result, ensure_ascii=False)
    assert "脱敏标题" not in serialized
    assert "脱敏正文" not in serialized
    assert working.read_bytes()
